from tqdm import tqdm, tqdm_gui
from docx import Document
import itertools, random, time, xlsxwriter, sys, copy, math

class Tournament():
    men=[]
    women=[]
    maxCount=[0,0,0,0,0,0,0,0,0]
    maxrepeats=500
    currentrepeats=0

    def __init__(self, nPlayers):
        self.men=["M"+str(player) for player in range(nPlayers)]
        print(list(self.men))
        self.women=["W"+str(player) for player in range(nPlayers)]
        print(list(self.women))  

    @staticmethod
    def GetTournamentRec(men, women, rounds, tableau):
        nPlayers=int(len(men))
        menVsMen=[[x for x in range(nPlayers)] for y in range(nPlayers)]
        menVsWomen=[[x for x in range(nPlayers)] for y in range(nPlayers)]
        menWithWomen=[[x for x in range(nPlayers)] for y in range(nPlayers)]
        womenVsWomen=[[x for x in range(nPlayers)] for y in range(nPlayers)]

        menIdx=[x for x in range(nPlayers)]
        womenIdx=[x for x in range(nPlayers)]
        IdxTableau=list()
        idxMenGames=list()
        idxWomenGames=list()
        IdxTableau.append(idxMenGames)
        IdxTableau.append(idxWomenGames)
      
        if Tournament.GetRounds(rounds, menIdx, womenIdx, IdxTableau, menVsMen, menVsWomen, menWithWomen, womenVsWomen):
            Tournament.IdxTableauToTableau(men, women, IdxTableau, tableau)
            return True
        else:
            return False
    @staticmethod
    def IdxTableauToTableau(men, women, idxTableau, tableau):
        print(idxTableau)
        menGames=[]
        womenGames=[] 
        for gameRoundIdx in range(int(len(idxTableau[0]))):
            
            menGame=[]
            womenGame=[]
            currentMenRound=idxTableau[0][gameRoundIdx]
            currentWomenRound=idxTableau[1][gameRoundIdx]       
            for idxGame in range(int(len(idxTableau[0][gameRoundIdx]))):
                currentMen=men[currentMenRound[idxGame]]
                currentWomen=women[currentWomenRound[idxGame]]
                menGame.append(currentMen)
                womenGame.append(currentWomen)
            menGames.append(menGame)
            womenGames.append(womenGame)
        tableau.append(menGames)
        tableau.append(womenGames)   
        

    @staticmethod
    def GetRounds(remainingrounds, men, women, tableau, menVsMen, menVsWomen, menWithWomen, womenVsWomen, fuzzyLimit=0):
        print("Remaining Rounds " + str(remainingrounds) + " Limit:" +str(fuzzyLimit))
        
        if remainingrounds<1:
            return True
        else:
            nPlayers=int(len(men))
            menGames=[-1 for x in range(nPlayers)]
            womenGames=[-1 for x in range(nPlayers)]
            menRec=copy.deepcopy(men)
            womenRec=copy.deepcopy(women)            
            if Tournament.GetRound(men=menRec, women=womenRec, menGames=menGames, womenGames=womenGames, menVsMen=menVsMen, menVsWomen=menVsWomen, menWithWomen=menWithWomen, womenVsWomen=womenVsWomen, GmenVsMen=menVsMen, GmenVsWomen=menVsWomen, GmenWithWomen=menWithWomen, GwomenVsWomen=womenVsWomen, fuzzyLimit=fuzzyLimit):
                tableau[0].append(menGames)
                tableau[1].append(womenGames)
                #TODO: when uneven - add remaining men/women to round? Or add intermediate round with remaining man (having double counter for remainingrounds): Remove Opponent or Partner for next round
                return Tournament.GetRounds(remainingrounds=remainingrounds-1,men=men, women=women, tableau=tableau, menVsMen=menVsMen, menVsWomen=menVsWomen, menWithWomen=menWithWomen, womenVsWomen=womenVsWomen)
            else:
                fuzzyLimit+=2
                input(">> Allowing double opponents")
                if fuzzyLimit<=nPlayers:
                    return Tournament.GetRounds(remainingrounds=remainingrounds,men=men, women=women, tableau=tableau, menVsMen=menVsMen, menVsWomen=menVsWomen, menWithWomen=menWithWomen, womenVsWomen=womenVsWomen, fuzzyLimit=fuzzyLimit)
                return False
               

    @staticmethod
    def GetRound(men, women, menGames, womenGames, menVsMen, menVsWomen, menWithWomen, womenVsWomen, GmenVsMen, GmenVsWomen, GmenWithWomen, GwomenVsWomen, fuzzyLimit=0):
        #TODO: Check for uneven, start with next round?
        if len(men)<2:
            tqdm.write("Termination, result found ")
            return True # termination of recursion
        else:
            
            Tournament.currentrepeats+=1
            hasResult=False
            localMenStart=men.copy()
            localWomen=women.copy()
            permCounter=0
            localMenPermutations=list()
            nPerm=1
            permIterator=localMenPermutations
            if len(localMenStart)<len(menVsWomen):
                localMenPermutations.append(localMenStart)               
            else:
                random.shuffle(localWomen)
                random.shuffle(localMenStart)
                localMenPermutations=itertools.permutations(localMenStart)
                nPerm=math.factorial(len(localMenStart))
                permIterator=tqdm(localMenPermutations, total=nPerm)

            for localMenTuple in permIterator:
                permCounter+=1
                localMen=list(localMenTuple)
                firstMan=localMen.pop(0)
                isFuzzyLimit=len(men)<=fuzzyLimit
                if isFuzzyLimit:
                    firstManVsMen=set(localMen)
                    firstManVsWomen=set(localWomen)
                else:
                    firstManVsMen=set(menVsMen[firstMan]) & set(localMen)
                    firstManVsWomen=set(menVsWomen[firstMan]) & set(localWomen)
                
                firstManWithWomen=set(menWithWomen[firstMan]) & set(localWomen)
                
                idxPlayer=int(len(men)/2-1)
                idxOpponent=int(len(men)/2-1)+int(len(menGames)/2)
            
                for opponent in firstManVsMen:
                    validRound=False
                    opponentWithWomen=set(menWithWomen[opponent])
                    
                    if isFuzzyLimit:
                        opponentVsWomen=set(localWomen)
                    else:
                        opponentVsWomen=set(menVsWomen[opponent])

                    firstManPartners=opponentVsWomen & firstManWithWomen & set(localWomen)
                    opponentPartners=opponentWithWomen & firstManVsWomen & set(localWomen)
                    
                    for firstManPartner in firstManPartners:
                        if isFuzzyLimit:
                            firstManPartnerOpponents=set(localWomen)
                        else:
                            firstManPartnerOpponents=set(womenVsWomen[firstManPartner]) & set(localWomen)

                        firstManPartnerSet={firstManPartner}
                        validOpponentPartners=(firstManPartnerOpponents & opponentPartners)-firstManPartnerSet
                       
                        validRound=len(validOpponentPartners) > 0
                        if validRound:
                            for opponentPartner in validOpponentPartners:
                                
                                if opponentPartner in womenVsWomen[firstManPartner]:
                                    localWomenRec=copy.deepcopy(women)
                                    localMenRec=copy.deepcopy(localMen)
                                    localMenVsMenRec=copy.deepcopy(menVsMen)
                                    localMenVsWomenRec=copy.deepcopy(menVsWomen)
                                    localMenWithWomenRec=copy.deepcopy(menWithWomen)
                                    localWomenVsWomenRec=copy.deepcopy(womenVsWomen) 
                                    Tournament.UpdateGameTrackers(localWomenRec, opponentPartner, firstManPartner, localMenRec, opponent, localMenVsMenRec, firstMan, localMenVsWomenRec, localMenWithWomenRec, localWomenVsWomenRec)
                                    hasResult=Tournament.GetRound(localMenRec, localWomenRec, menGames, womenGames, localMenVsMenRec, localMenVsWomenRec, localMenWithWomenRec, localWomenVsWomenRec, GmenVsMen, GmenVsWomen, GmenWithWomen, GwomenVsWomen, fuzzyLimit=fuzzyLimit)
                                    if hasResult:
                                        Tournament.UpdateGameTrackers(women, opponentPartner, firstManPartner, men, opponent, GmenVsMen, firstMan, GmenVsWomen, GmenWithWomen, GwomenVsWomen)
                                        menGames[idxPlayer]=firstMan
                                        menGames[idxOpponent]=opponent
                                        womenGames[idxPlayer]=firstManPartner
                                        womenGames[idxOpponent]=opponentPartner
                                        return True
                                    
            if hasResult:
                raise ValueError("result cannot be true!")
            else:
                return False


        raise ReferenceError("no valid branch followed in GetRound method")
    
    @staticmethod
    def UpdateGameTrackers(women, opponentPartner, firstManPartner, men, opponent, menVsMen, firstMan, menVsWomen, menWithWomen, womenVsWomen):
        #print(" *** UpdateGameTracker ***")
        #print("opponentPartner: " + str(opponentPartner))
        #print("opponent: " + str(opponent))
        #print("firstManPartner: "+str(firstManPartner))
        #print("firstMan: " +str(firstMan))
        #print("-womenbefore: " +str(women))
        limit=0

        women.remove(opponentPartner)
        women.remove(firstManPartner)
        #print("-womenafter: "+str(women))
        #print("-menbefore: "+str(men))
        if firstMan in men:
            men.remove(firstMan)
        men.remove(opponent)
        #print("-menafter: " +str(men))

        if opponent in menVsMen[firstMan]: menVsMen[firstMan].remove(opponent)
        Tournament.recover(menVsMen, firstMan, limit)
        if firstMan in menVsMen[opponent]: menVsMen[opponent].remove(firstMan)
        Tournament.recover(menVsMen, opponent, limit)
        #print("menVsWomen before opponentPartnerRemove: "  +str(menVsWomen[firstMan]))
        if opponentPartner in menVsWomen[firstMan]: menVsWomen[firstMan].remove(opponentPartner)
        Tournament.recover(menVsWomen, firstMan, limit)
        #print("menVsWomen after opponentParterRemove: " +str(menVsWomen[firstMan]))
        #print("menVsWomen before firstManPartnerRemove: "+str(menVsWomen[opponent]))
        if firstManPartner in menVsWomen[opponent]: menVsWomen[opponent].remove(firstManPartner)
        Tournament.recover(menVsWomen, opponent, limit)
        #print("menVsWomen after firstManParterRemove: " +str(menVsWomen[opponent]))

        #print("menwithWomen before firstManPartner removal: " +str(menWithWomen[firstMan]))
        if firstManPartner in menWithWomen[firstMan]: menWithWomen[firstMan].remove(firstManPartner)
        Tournament.recover(menWithWomen, firstMan, limit)
        #print("menwithWomen after firstManPartner removal: " +str(menWithWomen[firstMan]))
        #print("menwithWomen before opponentpartner
        #removal: " +str(menWithWomen[opponent]))
        if opponentPartner in menWithWomen[opponent]: menWithWomen[opponent].remove(opponentPartner)
        Tournament.recover(menWithWomen, opponent, limit)
        #print("menwithWomen after opponentpartner removal: " +str(menWithWomen[opponent]))

        #print("womenVsWomen before opponentPartner removal: " +str(womenVsWomen[firstManPartner]))
        if opponentPartner in womenVsWomen[firstManPartner]: womenVsWomen[firstManPartner].remove(opponentPartner)
        Tournament.recover(womenVsWomen, firstManPartner, limit)
        #print("womenVsWomen after opponentPartner removal: " +str(womenVsWomen[firstManPartner]))
        #print("womenVsWomen before firstManPartner removal: " +str(womenVsWomen[opponentPartner]))
        if firstManPartner in womenVsWomen[opponentPartner]: womenVsWomen[opponentPartner].remove(firstManPartner)
        Tournament.recover(womenVsWomen, opponentPartner, limit)   
        #print("womenVsWomen after firstManPartner removal: " +str(womenVsWomen[opponentPartner]))    

    @staticmethod
    def recover(listinput, index, limit):
        if len(listinput[index])<=limit:
            x=random.randint(0, len(listinput[index]))
            while x in listinput[index]:
                x=random.randint(0, len(listinput[index]))
            listinput[index].append(x)

    @staticmethod
    def GetOpponent(player, players, playedAgainst):
        return 0
    
    @staticmethod
    def GetPartner1(player1, player2, partners, playedAgainst):
        return 0
    
    @staticmethod
    def GetPartner2(playerG1, playerG2, opponentG1, opponents):
        return 0

    @staticmethod
    def NotPlayedTogether(player1, player2, partners):
        if partners[player1][player2]:
            return False
        else:
            return True
    
    @staticmethod
    def NotPlayedAgainst(player1, player2, opponents):
        if opponents[player1][player2]:
            return False
        else:
            return True
    
    @staticmethod
    def GetTournament(men, women, rounds, tableau):
        #return Tournament.GetTournamentBrute(men, women, rounds, tableau)
        return Tournament.GetTournamentRec(men, women, rounds, tableau)

    @staticmethod
    def GetTournamentBrute(men, women, rounds, tableau):
        foundResult=False
        mensGames=[]
        womensGames=[]
        print("Generating men...")
        Tournament.GetGames(men, mensGames, len(men)-1)
        print("Generated " + str(len(mensGames)) + " mens rounds")
        print("Generating women...")
        Tournament.GetGames(women, womensGames, len(women)-1)
        print ("Generated " + str (len(womensGames)) + " womens rounds")
        countAttempts=1
        while Tournament.ValidateGames(mensGames, womensGames, men, women, rounds)==False:
            womensGames.clear()
            print("Attempts:" + str(countAttempts) + " Results : " + str(list(Tournament.maxCount)))
            Tournament.GetGames(women, womensGames, len(women)-1)
            womensGamesPermut=list(itertools.permutations(womensGames))
            for aWomensGames in womensGamesPermut:
                womensGames.clear()
                womensGames=list(aWomensGames)
                print("Attempts:" + str(countAttempts) + " Results : " + str(list(Tournament.maxCount)))
                countAttempts+=1
                validationResult=Tournament.ValidateGames(mensGames, womensGames, men, women, rounds)
                if validationResult:
                    print("\nfound result!")
                    foundResult=True
                    break

            if countAttempts>sys.maxsize:
                print("No Tournament Found after " + str(countAttempts) + " attempts")
                break
            if foundResult==False:
                countAttempts+=1
                womensGames.clear()
                print("\nTaking " + str(len(mensGames)) + " mens rounds")
                print("Generating new womens rounds")
                Tournament.GetGames(women, womensGames, len(women)-1)
                print ("Generated " + str (len(womensGames)) + " new womens rounds")
            else:
                break

        if Tournament.ValidateGames(mensGames, womensGames, men, women, rounds)==True:
            print("SUCCESS")
            #Tournament.printGameTableau(mensGames, womensGames, rounds)
        tableau.clear()
        tableau.append(mensGames)
        tableau.append(womensGames)
    
    @staticmethod
    def printGameTableau(mensGames, womensGames, rounds):
        print("-------------------------------------------------------")
        for x in range(rounds):
            playersMen=mensGames[x]
            playersWomen=womensGames[x]
            for p in range(int(len(playersMen)/2)):
                player1=str(playersMen[p])
                partner1=str(playersWomen[p])
                player2=str(playersMen[int(len(playersMen)/2)+p])
                partner2=str(playersWomen[int(len(playersMen)/2)+p])
                print("Runde " + str(x+1)+ ": " + player1 + " & " + partner1 + " vs. " + player2 + " & " + partner2)
            print("-------------------------------------------------------")
    
    @staticmethod
    def ValidateGames(tableauA, tableauB, men, women, rounds):
        success=True
        if len(tableauA)==len(tableauB):
            nPlayers=len(men)
            nGames=int(nPlayers/2)
            partners=[[False for x in range(nPlayers)] for y in range(nPlayers)]
            opponents=[[False for x in range(nPlayers)] for y in range(nPlayers)]
            womenOpponents=[[False for x in range(nPlayers)] for y in range(nPlayers)]  
            maleOpponents=[[False for x in range(nPlayers)] for y in range(nPlayers)]
                
            for gamesA in tableauA:
                roundIndex=tableauA.index(gamesA)
                gamesB=tableauB[roundIndex]
                # print(str(roundIndex+1) + " / " + str(list(gamesA))+ " / " + str(list(gamesB)))
                for playerA in gamesA:
                    gameIndex=gamesA.index(playerA)
                    playerB=gamesB[gameIndex]
                    playerAindex=men.index(playerA)
                    playerBindex=women.index(playerB)
                    if partners[playerAindex][playerBindex]==True:
                        success=False
                        Tournament.maxCount[roundIndex]+=1
                        if (roundIndex+1)>=rounds:
                            success=True
                        break
                    else:
                        partners[playerAindex][playerBindex]=True

                if success==False:
                    break
                else:
                    for gameID in range(nGames):
                        maleOpponent1=gamesA[gameID]
                        maleOpponent2=gamesA[gameID+nGames]
                        womenOpponent1=gamesB[gameID]
                        womenOpponent2=gamesB[gameID+nGames]
                        maleOpponent1Index=men.index(maleOpponent1)
                        maleOpponent2Index=men.index(maleOpponent2)
                        womenOpponent1Index=women.index(womenOpponent1)
                        womenOpponent2Index=women.index(womenOpponent2)
                        if (maleOpponents[maleOpponent1Index][maleOpponent2Index]==True) \
                            or (womenOpponents[womenOpponent1Index][womenOpponent2Index]==True) \
                            or (maleOpponents[maleOpponent2Index][maleOpponent1Index]==True) \
                            or (womenOpponents[womenOpponent2Index][womenOpponent1Index]==True) \
                            or (opponents[maleOpponent1Index][womenOpponent2Index]==True) \
                            or (opponents[maleOpponent2Index][womenOpponent1Index]==True):

                            success=False
                            Tournament.maxCount[roundIndex]+=1
                            if (roundIndex+1)>=rounds:
                                success=True
                            break
                        else:
                            opponents[maleOpponent1Index][womenOpponent2Index]=True
                            opponents[maleOpponent2Index][womenOpponent1Index]=True
                            maleOpponents[maleOpponent1Index][maleOpponent2Index]=True
                            maleOpponents[maleOpponent2Index][maleOpponent1Index]=True
                            womenOpponents[womenOpponent1Index][womenOpponent2Index]=True
                            womenOpponents[womenOpponent2Index][womenOpponent1Index]=True

                    if success==False:
                        break

        else:
            success=False

        return success   

    @staticmethod
    def GetGames(players, games, rounds):
        if players:
            nPlayers=len(players)
            
            opponents=[[0 for x in range(nPlayers)] for y in range(nPlayers)] 
            for op1 in range(nPlayers):
                for op2 in range(nPlayers):
                    opponents[op1][op2]=False
            
            allGames=list(itertools.permutations(players))
            random.shuffle(allGames)
       
            for tournamentRound in allGames:
                success=True
                for game in range(int(nPlayers/2)):
                    player1=tournamentRound[game]
                    player2=tournamentRound[game+int(nPlayers/2)]
                    posPlayer1=players.index(player1)
                    posPlayer2=players.index(player2)
                    if (opponents[posPlayer1][posPlayer2]==True) or (opponents[posPlayer2][posPlayer1]==True):
                        success=False
                        break
                        
                if success==True:
                    games.append(tournamentRound)
                    for game in range(int(nPlayers/2)):
                        player1=tournamentRound[game]
                        player2=tournamentRound[game+int(nPlayers/2)]
                        posPlayer1=players.index(player1)
                        posPlayer2=players.index(player2)
                        opponents[posPlayer1][posPlayer2]=True
                        opponents[posPlayer2][posPlayer1]=True

                if len(games)>=rounds:
                    break
            if(len(games)<rounds):
                games.clear()
                Tournament.GetGames(players, games, rounds)
        else:
            print("No Players found")

    @staticmethod
    def WriteWord(fileName, nGroups, malePlayers, womenPlayers, mensGames, womensGames, nParallel):
        print("-> Writing Word File")
        doc=Document()
        doc.add_heading("Mixed Turnier", level=1)
        doc.add_heading("Teilnehmer", level=2)
        
        for aGrp in range(nGroups):
            doc.add_heading("Gruppe " + str(aGrp+1), level=3)
            doc.add_heading("Männer", level=4)
            nRows=len(malePlayers)+1
            nCols=2+len(mensGames)
            table=doc.add_table(rows=nRows, cols=nCols)
            table.rows[0].cells[0].text="Spieler"
            table.rows[0].cells[1].text="Name"
            for aGame in range(len(mensGames)):
                table.rows[0].cells[2+aGame].text=str(aGame+1)

            for aPlayer in range(len(malePlayers)):
                table.rows[1+aPlayer].cells[0].text=malePlayers[aPlayer]
            doc.add_heading("Frauen", level=4)
            nRows=len(womenPlayers)+1
            nCols=2+len(womensGames)
            table=doc.add_table(rows=nRows, cols=nCols)
            table.rows[0].cells[0].text="Spielerin"
            table.rows[0].cells[1].text="Name"
            for aGame in range(len(mensGames)):
                table.rows[0].cells[2+aGame].text=str(aGame+1)

            for aPlayer in range(len(malePlayers)):
                table.rows[1+aPlayer].cells[0].text=malePlayers[aPlayer]
        
        doc.save(fileName+".docx")
                
    @staticmethod
    def WriteExcel(fileName, nGroups, malePlayers, womenPlayers, mensGames, womensGames, nParallel):
        print("-> Writing Excel File")
        wb=xlsxwriter.Workbook(fileName+".xlsx")
        playersSheet=wb.add_worksheet("Spieler")
        tableauSheet=wb.add_worksheet("Turnier")
        nGamesPerRound=int(len(malePlayers)/2)
        nParallelGamesPerGroup=int(nParallel/2)
        nGames=int(len(mensGames))
        currentRow=1
        for aGroup in range(nGroups):
            playersSheet.write(currentRow, 1, "Gruppe " + str(aGroup+1))
            currentRow+=1
            for aPlayer in range(len(malePlayers)):
                maleFormula="="
                womenFormula="="
                maleFormula2="="
                womenFormula2="="

                for aRound in range(len(mensGames)):
                    maleFormula+="+G"+str(aGroup+1)+"_"+malePlayers[aPlayer]+"_R"+str(aRound)+"_Games"
                    womenFormula+="+G"+str(aGroup+1)+"_"+womenPlayers[aPlayer]+"_R"+str(aRound)+"_Games"
                    maleFormula2+="+G"+str(aGroup+1)+"_"+malePlayers[aPlayer]+"_R"+str(aRound)+"_Wins"
                    womenFormula+="+G"+str(aGroup+1)+"_"+womenPlayers[aPlayer]+"_R"+str(aRound)+"_Wins"

                maleCell=xlsxwriter.utility.xl_rowcol_to_cell(currentRow,2,col_abs=True, row_abs=True)
                playersSheet.write(maleCell, malePlayers[aPlayer])
                playersSheet.write(currentRow, 3, maleFormula)
                playersSheet.write(currentRow, 4, maleFormula2)
                womenCell=xlsxwriter.utility.xl_rowcol_to_cell(currentRow,7, col_abs=True, row_abs=True)
                wb.define_name("G"+str(aGroup+1)+"_"+malePlayers[aPlayer],"="+playersSheet.name+"!"+maleCell)
                playersSheet.write(womenCell, womenPlayers[aPlayer])
                playersSheet.write(currentRow, 8, womenFormula)
                playersSheet.write(currentRow, 9, womenFormula2)
                wb.define_name("G"+str(aGroup+1)+"_"+womenPlayers[aPlayer],"="+playersSheet.name+"!"+womenCell)
                currentRow+=1
            currentRow+=1
        currentRow=1
        currentGame=0
        for aRound in range(len(mensGames)):
            currentGamesMen=mensGames[aRound]
            currentGamesWomen=womensGames[aRound]
            for aGame in range(nGamesPerRound):
                man1=currentGamesMen[aGame]
                man2=currentGamesMen[aGame+nGamesPerRound]
                women1=currentGamesWomen[aGame]
                women2=currentGamesWomen[aGame+nGamesPerRound]
               
                for aGroup in range(nGroups):
                    tableauSheet.write(currentRow+aGroup*nParallelGamesPerGroup, 1, "Runde " +str(aRound+1))
                    tableauSheet.write(currentRow+aGroup*nParallelGamesPerGroup, 3, "Gruppe "+str(aGroup+1))
                    tableauSheet.write(currentRow+aGroup*nParallelGamesPerGroup, 4, "Spiel "+str(currentGame+1))
                    tableauSheet.write(currentRow+aGroup*nParallelGamesPerGroup, 5, "="+"G"+str(aGroup+1)+"_"+man1)
                    tableauSheet.write(currentRow+aGroup*nParallelGamesPerGroup, 6,"=G"+str(aGroup+1)+"_"+ women1)
                    result1LeftCell=xlsxwriter.utility.xl_rowcol_to_cell(currentRow+aGroup*nParallelGamesPerGroup,7, col_abs=True, row_abs=True)
                    result2LeftCell=xlsxwriter.utility.xl_rowcol_to_cell(currentRow+aGroup*nParallelGamesPerGroup,8, col_abs=True, row_abs=True)
                    wb.define_name("G"+str(aGroup+1)+"_"+man1+"_R"+str(aRound)+"_Games","="+tableauSheet.name+"!"+result1LeftCell)
                    wb.define_name("G"+str(aGroup+1)+"_"+women1+"_R"+str(aRound)+"_Games","="+tableauSheet.name+"!"+result1LeftCell)
                    wb.define_name("G"+str(aGroup+1)+"_"+man1+"_R"+str(aRound)+"_Wins","="+tableauSheet.name+"!"+result2LeftCell)
                    wb.define_name("G"+str(aGroup+1)+"_"+women1+"_R"+str(aRound)+"_Wins","="+tableauSheet.name+"!"+result2LeftCell)
                    tableauSheet.write(currentRow+aGroup*nParallelGamesPerGroup, 9, "gegen")
                    tableauSheet.write(currentRow+aGroup*nParallelGamesPerGroup, 10, "="+"G"+str(aGroup+1)+"_"+man2)
                    tableauSheet.write(currentRow+aGroup*nParallelGamesPerGroup, 11,"="+"G"+str(aGroup+1)+"_"+ women2)
                    result1RightCell=xlsxwriter.utility.xl_rowcol_to_cell(currentRow+aGroup*nParallelGamesPerGroup,12, col_abs=True, row_abs=True)
                    result2RightCell=xlsxwriter.utility.xl_rowcol_to_cell(currentRow+aGroup*nParallelGamesPerGroup,13, col_abs=True, row_abs=True)
                    wb.define_name("G"+str(aGroup+1)+"_"+man2+"_R"+str(aRound)+"_Games","="+tableauSheet.name+"!"+result1RightCell)
                    wb.define_name("G"+str(aGroup+1)+"_"+women2+"_R"+str(aRound)+"_Games","="+tableauSheet.name+"!"+result1RightCell)
                    wb.define_name("G"+str(aGroup+1)+"_"+man2+"_R"+str(aRound)+"_Wins","="+tableauSheet.name+"!"+result2RightCell)
                    wb.define_name("G"+str(aGroup+1)+"_"+women2+"_R"+str(aRound)+"_Wins","="+tableauSheet.name+"!"+result2RightCell)
                    tableauSheet.write_formula(currentRow+aGroup*nParallelGamesPerGroup,8, "=if("+result1LeftCell+">"+result1RightCell+",1,0)")
                    tableauSheet.write_formula(currentRow+aGroup*nParallelGamesPerGroup,13, "=if("+result1RightCell+">"+result1LeftCell+",1,0)")

                currentGame+=1
                if currentGame % nParallelGamesPerGroup==0:
                    currentRow+=nGroups*nParallelGamesPerGroup
                else:
                   currentRow+=1

        wb.close()

    @staticmethod
    def BatchRunner():
        print("*****************************")
        print("*** Mixed Tournament v2.0 ***")
        print("*****************************")
        tournamentResult=[]
                       
        print("Current Batch: 10/6/2")
        Tournament.CreateTournament(nPlayers=10, nRounds=6, tournamentResult=tournamentResult, start=time.time(), filename="2020_10_6_2", nGroups=2, nParallel=4)
        print("Current Batch: 11/6/2")
        Tournament.CreateTournament(nPlayers=11, nRounds=6, tournamentResult=tournamentResult, start=time.time(), filename="2020_11_6_2", nGroups=2, nParallel=4)
        print("Current Batch: 12/6/2")
        Tournament.CreateTournament(nPlayers=12, nRounds=6, tournamentResult=tournamentResult, start=time.time(), filename="2020_12_6_2", nGroups=2, nParallel=4)
        print("Current Batch: 13/6/2")
        Tournament.CreateTournament(nPlayers=13, nRounds=6, tournamentResult=tournamentResult, start=time.time(), filename="2020_13_6_2", nGroups=2, nParallel=4)
        print("Current Batch: 20/6/1")
        Tournament.CreateTournament(nPlayers=20, nRounds=6, tournamentResult=tournamentResult, start=time.time(), filename="2020_20_6_1", nGroups=1, nParallel=4)
        print("Current Batch: 21/6/1")
        Tournament.CreateTournament(nPlayers=21, nRounds=6, tournamentResult=tournamentResult, start=time.time(), filename="2020_21_6_1", nGroups=1, nParallel=4)
        print("Current Batch: 22/6/1")
        Tournament.CreateTournament(nPlayers=22, nRounds=6, tournamentResult=tournamentResult, start=time.time(), filename="2020_22_6_1", nGroups=1, nParallel=4)
        print("Current Batch: 23/6/1")
        Tournament.CreateTournament(nPlayers=23, nRounds=6, tournamentResult=tournamentResult, start=time.time(), filename="2020_23_6_1", nGroups=1, nParallel=4)
        print("Current Batch: 24/6/1")
        Tournament.CreateTournament(nPlayers=24, nRounds=6, tournamentResult=tournamentResult, start=time.time(), filename="2020_24_6_1", nGroups=1, nParallel=4)
        print("Current Batch: 25/6/1")
        Tournament.CreateTournament(nPlayers=25, nRounds=6, tournamentResult=tournamentResult, start=time.time(), filename="2020_25_6_1", nGroups=1, nParallel=4)
        print("Current Batch: 26/6/1")
        Tournament.CreateTournament(nPlayers=26, nRounds=6, tournamentResult=tournamentResult, start=time.time(), filename="2020_26_6_1", nGroups=1, nParallel=4)
        print("Current Batch: 27/6/1")
        Tournament.CreateTournament(nPlayers=27, nRounds=6, tournamentResult=tournamentResult, start=time.time(), filename="2020_27_6_1", nGroups=1, nParallel=4)

   
    @staticmethod
    def Runner():
        start=time.time()
        print("*****************************")
        print("*** Mixed Tournament v2.0 ***")
        print("*****************************")
        tournamentResult=[]
        nPlayers=int(input("Number of players: "))
        nRounds=int(input("Number of rounds: "))
        nGroups=int(input("Number of tournament groups: "))
        nParallel=int(input("Number of parallel games: "))
        filename=str(input("Filename for output: "))
        Tournament.CreateTournament(nPlayers, nRounds, tournamentResult, start, filename, nGroups, nParallel)

    @staticmethod
    def CreateTournament(nPlayers, nRounds, tournamentResult, start, filename, nGroups, nParallel):
        Tournament.maxCount=[0,0,0,0,0,0,0,0,0]
        tournamentResult.clear()
        malePlayers=["M"+str(player) for player in range(nPlayers)]
        womenPlayers=["W"+str(player) for player in range(nPlayers)]
        print("-> starting calc...")
        hadSuccess=Tournament.GetTournament(malePlayers, womenPlayers, nRounds, tournamentResult)
        end=time.time()
        duration=end-start
        print("-> finished calc after " + str(duration)+ " seconds")
        if hadSuccess:
            print("\nFinal Result:")
            Tournament.printGameTableau(tournamentResult[0], tournamentResult[1], nRounds)
            Tournament.WriteExcel(filename, nGroups, malePlayers, womenPlayers,tournamentResult[0][:nRounds], tournamentResult[1][:nRounds], nParallel)
            Tournament.WriteWord(filename, nGroups, malePlayers, womenPlayers,tournamentResult[0][:nRounds], tournamentResult[1][:nRounds], nParallel)
        else:
            print("\n*** no result found ***")

# EXECUTION STARTS HERE
Tournament.Runner()
